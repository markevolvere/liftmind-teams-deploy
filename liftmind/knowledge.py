"""
LiftMind Knowledge Extraction Engine

Multi-layer document processing:
1. Document metadata extraction
2. Smart semantic chunking
3. Fact extraction using Claude
4. Q&A pair generation
5. Entity extraction and registry
"""

import os
import re
import json
import logging
import hashlib
from typing import Optional, List, Dict, Any
from pathlib import Path
from datetime import datetime

import psycopg2
from psycopg2.extras import Json, execute_values
import pdfplumber

from liftmind.config import settings
from liftmind.embedding_utils import generate_embedding, generate_embeddings_batch

logger = logging.getLogger(__name__)

# ============================================================================
# TOKEN COUNTING
# ============================================================================

# Try to use tiktoken for accurate token counting, fall back to estimation
_tiktoken_encoder = None

def _init_tiktoken():
    """Initialize tiktoken encoder (lazy loading)."""
    global _tiktoken_encoder
    if _tiktoken_encoder is not None:
        return _tiktoken_encoder

    try:
        import tiktoken
        # Use cl100k_base encoding (used by GPT-4, Claude-compatible token counts)
        _tiktoken_encoder = tiktoken.get_encoding("cl100k_base")
        logger.info("tiktoken initialized for accurate token counting")
    except ImportError:
        logger.warning("tiktoken not installed - using word count estimation for tokens. "
                      "Install with: pip install tiktoken")
        _tiktoken_encoder = False  # Mark as unavailable
    except Exception as e:
        logger.warning(f"tiktoken initialization failed: {e} - using word count estimation")
        _tiktoken_encoder = False

    return _tiktoken_encoder


def count_tokens(text: str) -> int:
    """
    Count tokens in text using tiktoken (accurate) or word estimation (fallback).

    Uses cl100k_base encoding which is compatible with modern LLMs like GPT-4 and Claude.
    Falls back to words * 1.3 estimation if tiktoken is unavailable.

    Args:
        text: Text to count tokens for

    Returns:
        Token count (accurate if tiktoken available, estimated otherwise)
    """
    encoder = _init_tiktoken()

    if encoder and encoder is not False:
        try:
            return len(encoder.encode(text))
        except Exception as e:
            logger.debug(f"tiktoken encoding failed: {e}, falling back to estimation")

    # Fallback: estimate tokens as words * 1.3
    return int(len(text.split()) * 1.3)

# ============================================================================
# DATABASE CONNECTION
# ============================================================================

def get_db():
    """
    Get database connection.

    NOTE: This function maintains backwards compatibility. For new code using
    context managers, prefer `from liftmind.database import get_connection`.
    """
    return psycopg2.connect(settings.DATABASE_URL)


def init_schema():
    """Initialize database schema."""
    schema_path = os.path.join(os.path.dirname(__file__), '..', 'scripts', 'schema.sql')

    conn = get_db()
    cur = conn.cursor()

    with open(schema_path, 'r') as f:
        # Execute schema, ignoring extension errors (might not have pg_trgm/vector)
        sql = f.read()
        try:
            cur.execute(sql)
        except Exception as e:
            # Rollback and try without vector extension
            conn.rollback()
            sql_no_vector = sql.replace("CREATE EXTENSION IF NOT EXISTS vector;", "-- vector extension not available")
            cur.execute(sql_no_vector)

    conn.commit()
    cur.close()
    conn.close()
    logger.info("Schema initialized")


# ============================================================================
# DOCUMENT REGISTRATION
# ============================================================================

SUPPORTED_EXTENSIONS = {
    'pdf': ['.pdf'],
    'document': ['.doc', '.docx'],
    'text': ['.txt', '.md'],
    'image': ['.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp']
}


def get_file_type(filepath: str) -> Optional[str]:
    """Determine file type from extension."""
    ext = Path(filepath).suffix.lower()
    for file_type, extensions in SUPPORTED_EXTENSIONS.items():
        if ext in extensions:
            return file_type
    return None


def compute_file_hash(filepath: str) -> str:
    """Compute SHA256 hash of file."""
    sha256 = hashlib.sha256()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            sha256.update(chunk)
    return sha256.hexdigest()


def detect_lift_models(text: str, filename: str) -> List[str]:
    """Detect lift models mentioned in text or filename."""
    # Known lift model patterns
    known_models = [
        'Elfo', 'Elfo 2', 'Elfo XL', 'Elfo Traction', 'Elfo Hydraulic',
        'Bari', 'E3', 'Tresa', 'Freedom', 'Freedom MAXI', 'Freedom STEP',
        'P4', 'Pollock', 'Supermec', 'Supermec 2', 'Supermec 3',
        'Elvoron', 'Boutique'
    ]

    found = set()
    text_lower = (text + ' ' + filename).lower()

    for model in known_models:
        if model.lower() in text_lower:
            found.add(model)

    return list(found)


def detect_doc_type(filename: str, content_sample: str = "") -> str:
    """Detect document type from filename and content."""
    filename_lower = filename.lower()
    content_lower = content_sample.lower()

    if any(x in filename_lower for x in ['wiring', 'electrical', 'circuit', 'diagram']):
        return 'wiring_diagram'
    if any(x in filename_lower for x in ['manual', 'instruction']):
        return 'manual'
    if 'bulletin' in filename_lower:
        return 'bulletin'
    if any(x in filename_lower for x in ['spec', 'specification']):
        return 'spec_sheet'
    if any(x in filename_lower for x in ['assembly', 'installation']):
        return 'installation_guide'
    if 'troubleshoot' in filename_lower or 'fault' in filename_lower:
        return 'troubleshooting'

    return 'document'


def register_document(filepath: str, lift_model_hint: str = None) -> int:
    """Register a document in the database and return its ID."""
    filename = os.path.basename(filepath)
    file_type = get_file_type(filepath)
    file_size = os.path.getsize(filepath)
    file_hash = compute_file_hash(filepath)

    # Detect lift models from folder structure or filename
    lift_models = []
    if lift_model_hint:
        lift_models = [lift_model_hint]
    else:
        lift_models = detect_lift_models(filename, "")

    doc_type = detect_doc_type(filename)

    conn = get_db()
    cur = conn.cursor()

    # Check if document exists and hasn't changed
    cur.execute("""
        SELECT id, file_hash FROM documents WHERE file_path = %s
    """, (filepath,))

    existing = cur.fetchone()
    if existing:
        if existing[1] == file_hash:
            # No changes, return existing ID
            cur.close()
            conn.close()
            return existing[0]
        else:
            # File changed, update and re-index
            cur.execute("""
                UPDATE documents SET
                    file_hash = %s,
                    file_size = %s,
                    index_status = 'pending',
                    updated_at = NOW()
                WHERE id = %s
                RETURNING id
            """, (file_hash, file_size, existing[0]))
            doc_id = cur.fetchone()[0]

            # Clear old data
            cur.execute("DELETE FROM chunks WHERE document_id = %s", (doc_id,))
            cur.execute("DELETE FROM facts WHERE document_id = %s", (doc_id,))
    else:
        # New document
        cur.execute("""
            INSERT INTO documents (filename, file_path, file_type, doc_type, lift_models, file_size, file_hash)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (filename, filepath, file_type, doc_type, lift_models, file_size, file_hash))
        doc_id = cur.fetchone()[0]

    conn.commit()
    cur.close()
    conn.close()

    return doc_id


# ============================================================================
# SMART CHUNKING
# ============================================================================

def _extract_page_numbers(text: str) -> List[int]:
    """
    Extract page numbers from page markers like '--- Page N ---'.

    D10: Helper function to parse page markers for page tracking.
    """
    page_pattern = r'---\s*Page\s+(\d+)\s*---'
    matches = re.findall(page_pattern, text)
    return [int(m) for m in matches]


def smart_chunk_text(text: str, max_tokens: int = 400, overlap_ratio: float = 0.3) -> List[Dict]:
    """
    Smart chunking that respects document structure.

    - Keeps tables intact
    - Keeps numbered steps together
    - Respects section boundaries
    - Flexible chunk size based on content
    - Tracks section navigation paths
    - D10: Tracks page_start and page_end per chunk
    """
    chunks = []

    # Track section hierarchy for navigation paths
    current_section_path = []

    # Split by major sections (headers) - use \n\n to avoid splitting mid-paragraph
    # The lookahead matches section headers while preserving content between them
    section_pattern = r'\n\n(?=[A-Z][A-Z\s]{2,}:|\d+\.\s+[A-Z]|={3,}|-{3,})'
    sections = re.split(section_pattern, text)
    # Ensure no content is silently dropped - if splitting produced nothing, use full text
    if not sections or (len(sections) == 1 and not sections[0].strip()):
        sections = [text]

    current_chunk = ""
    current_tokens = 0
    current_chunk_section_path = ""
    # D10: Track pages covered by current chunk
    current_chunk_pages = []

    for section in sections:
        section = section.strip()
        if not section:
            continue

        # D10: Extract page numbers from this section
        section_pages = _extract_page_numbers(section)

        # Detect section headers and update path
        header_match = re.match(r'^([A-Z][A-Z\s]{2,}):?|^(\d+\.)\s+([A-Z].+?)(?:\n|$)', section)
        if header_match:
            header_text = header_match.group(1) or f"{header_match.group(2)} {header_match.group(3) or ''}"
            header_text = header_text.strip()

            # Determine header level (numbered = subsection, CAPS = main section)
            if header_match.group(2):  # Numbered section like "1. INSTALLATION"
                # Sub-section - add to path
                if len(current_section_path) > 1:
                    current_section_path = current_section_path[:1]  # Keep main section
                current_section_path.append(header_text)
            else:
                # Main section - reset path
                current_section_path = [header_text]

        # Build current section path string
        section_path_str = " > ".join(current_section_path) if current_section_path else ""

        # Accurate token count using tiktoken (or fallback estimation)
        section_tokens = count_tokens(section)

        # If section fits in current chunk, add it
        if current_tokens + section_tokens <= max_tokens:
            current_chunk += "\n\n" + section if current_chunk else section
            current_tokens += section_tokens
            if not current_chunk_section_path:
                current_chunk_section_path = section_path_str
            # D10: Track pages
            current_chunk_pages.extend(section_pages)
        else:
            # Save current chunk if it has content
            if current_chunk.strip():
                # D10: Calculate page_start and page_end
                page_start = min(current_chunk_pages) if current_chunk_pages else None
                page_end = max(current_chunk_pages) if current_chunk_pages else None
                chunks.append({
                    "content": current_chunk.strip(),
                    "tokens": current_tokens,
                    "section_path": current_chunk_section_path,
                    "page_start": page_start,
                    "page_end": page_end
                })

            # If section is too large, split it further
            if section_tokens > max_tokens:
                sub_chunks = _split_large_section(section, max_tokens, overlap_ratio)
                # Add section path and page info to sub-chunks
                for sub_chunk in sub_chunks:
                    sub_chunk["section_path"] = section_path_str
                    # D10: Parse pages from sub-chunk content
                    sub_pages = _extract_page_numbers(sub_chunk["content"])
                    sub_chunk["page_start"] = min(sub_pages) if sub_pages else (min(section_pages) if section_pages else None)
                    sub_chunk["page_end"] = max(sub_pages) if sub_pages else (max(section_pages) if section_pages else None)
                chunks.extend(sub_chunks)
                current_chunk = ""
                current_tokens = 0
                current_chunk_section_path = ""
                current_chunk_pages = []
            else:
                current_chunk = section
                current_tokens = section_tokens
                current_chunk_section_path = section_path_str
                current_chunk_pages = section_pages.copy()

    # Don't forget the last chunk
    if current_chunk.strip():
        # D10: Calculate page_start and page_end for final chunk
        page_start = min(current_chunk_pages) if current_chunk_pages else None
        page_end = max(current_chunk_pages) if current_chunk_pages else None
        chunks.append({
            "content": current_chunk.strip(),
            "tokens": current_tokens,
            "section_path": current_chunk_section_path,
            "page_start": page_start,
            "page_end": page_end
        })

    # Add overlap between chunks
    chunks = _add_overlap(chunks, overlap_ratio)

    return chunks


def _split_large_section(text: str, max_tokens: int, overlap_ratio: float) -> List[Dict]:
    """Split a large section while keeping logical units together."""
    chunks = []

    # Try to split by paragraphs first
    paragraphs = text.split('\n\n')

    current_chunk = ""
    current_tokens = 0

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # Use accurate token counting
        para_tokens = count_tokens(para)

        if current_tokens + para_tokens <= max_tokens:
            current_chunk += "\n\n" + para if current_chunk else para
            current_tokens += para_tokens
        else:
            if current_chunk.strip():
                chunks.append({
                    "content": current_chunk.strip(),
                    "tokens": current_tokens
                })

            # If paragraph itself is too large, do simple split
            if para_tokens > max_tokens:
                words = para.split()
                chunk_words = []
                chunk_token_count = 0
                for word in words:
                    chunk_words.append(word)
                    # Re-count periodically to maintain accuracy
                    if len(chunk_words) % 20 == 0:
                        chunk_token_count = count_tokens(" ".join(chunk_words))
                    else:
                        chunk_token_count += 1  # Rough increment

                    if chunk_token_count >= max_tokens * 0.9:
                        chunk_text = " ".join(chunk_words)
                        chunks.append({
                            "content": chunk_text,
                            "tokens": count_tokens(chunk_text)
                        })
                        chunk_words = []
                        chunk_token_count = 0
                if chunk_words:
                    current_chunk = " ".join(chunk_words)
                    current_tokens = count_tokens(current_chunk)
                else:
                    current_chunk = ""
                    current_tokens = 0
            else:
                current_chunk = para
                current_tokens = para_tokens

    if current_chunk.strip():
        chunks.append({
            "content": current_chunk.strip(),
            "tokens": count_tokens(current_chunk.strip())
        })

    return chunks


def _add_overlap(chunks: List[Dict], overlap_ratio: float) -> List[Dict]:
    """Add overlapping context between chunks, preserving metadata."""
    if len(chunks) <= 1:
        return chunks

    enhanced_chunks = []
    for i, chunk in enumerate(chunks):
        content = chunk["content"]

        # Add context from previous chunk
        if i > 0:
            prev_content = chunks[i-1]["content"]
            overlap_words = int(len(prev_content.split()) * overlap_ratio)
            if overlap_words > 10:
                overlap_text = " ".join(prev_content.split()[-overlap_words:])
                content = f"[...{overlap_text}]\n\n{content}"

        enhanced_chunk = {
            "content": content,
            "tokens": int(len(content.split()) * 1.3)
        }
        # Preserve section_path and other metadata from original chunk
        if "section_path" in chunk:
            enhanced_chunk["section_path"] = chunk["section_path"]
        if "chunk_type" in chunk:
            enhanced_chunk["chunk_type"] = chunk["chunk_type"]
        if "page_number" in chunk:
            enhanced_chunk["page_number"] = chunk["page_number"]
        # D10: Preserve page_start and page_end
        if "page_start" in chunk:
            enhanced_chunk["page_start"] = chunk["page_start"]
        if "page_end" in chunk:
            enhanced_chunk["page_end"] = chunk["page_end"]

        enhanced_chunks.append(enhanced_chunk)

    return enhanced_chunks


def detect_chunk_type(content: str) -> str:
    """Detect what type of content a chunk contains."""
    content_lower = content.lower()

    # Table detection
    if '|' in content and content.count('|') > 3:
        return 'table'

    # Numbered steps
    if re.search(r'^\s*\d+\.\s', content, re.MULTILINE):
        return 'procedure'

    # Warning/caution
    if any(x in content_lower for x in ['warning', 'caution', 'danger', '\u26a0']):
        return 'warning'

    # Bullet list
    if re.search(r'^\s*[-\u2022]\s', content, re.MULTILINE):
        return 'list'

    # Specifications (lots of numbers and units)
    if re.search(r'\d+\s*(mm|cm|m|kg|N|Nm|V|A|Hz|\u00b0|ohm|\u03a9)', content):
        return 'specifications'

    return 'text'


# ============================================================================
# TEXT EXTRACTION
# ============================================================================

def _extract_tables_as_markdown(page, page_num: int) -> List[Dict]:
    """
    Extract tables from a pdfplumber page and convert to markdown format.
    Repeats header every 15 rows for large tables.
    """
    table_chunks = []
    tables = page.extract_tables()

    for table in tables:
        if not table or len(table) < 2:
            continue

        # First row is header
        header = table[0]
        header_row = "| " + " | ".join(str(h or "").strip() for h in header) + " |"
        separator = "| " + " | ".join("---" for _ in header) + " |"

        # Build data rows
        data_rows = []
        for row in table[1:]:
            row_text = "| " + " | ".join(str(cell or "").strip() for cell in row) + " |"
            data_rows.append(row_text)

        if not data_rows:
            continue

        # Build full table markdown
        full_table = "\n".join([header_row, separator] + data_rows)

        # Large table? Split by rows but REPEAT header every 15 rows
        if len(full_table) > 3000 or len(data_rows) > 15:
            for i in range(0, len(data_rows), 15):
                chunk_rows = [header_row, separator] + data_rows[i:i+15]
                table_chunks.append({
                    "content": "\n".join(chunk_rows),
                    "chunk_type": "table",
                    "page_number": page_num
                })
        else:
            table_chunks.append({
                "content": full_table,
                "chunk_type": "table",
                "page_number": page_num
            })

    return table_chunks


def extract_pdf_text(filepath: str) -> tuple[str, int, Dict, List[Dict]]:
    """
    Extract text and structure from PDF using pdfplumber for better table extraction.
    Falls back to PyPDF if pdfplumber fails.

    Returns:
        tuple: (full_text, page_count, toc, table_chunks)
        - table_chunks have explicit chunk_type='table' set during extraction
    """
    try:
        return _extract_pdf_with_pdfplumber(filepath)
    except Exception as e:
        logger.warning(f"pdfplumber failed for {filepath}, falling back to PyPDF: {e}")
        return _extract_pdf_with_pypdf(filepath)


def _extract_pdf_with_pdfplumber(filepath: str) -> tuple[str, int, Dict, List[Dict]]:
    """
    Extract text and tables from PDF using pdfplumber.

    Returns:
        tuple: (full_text, page_count, toc, table_chunks)
        - table_chunks are returned separately with explicit chunk_type='table'

    Uses batch processing for large PDFs (100+ pages) to avoid memory issues.
    """
    full_text = ""
    toc = {"pages": []}
    table_chunks = []

    with pdfplumber.open(filepath) as pdf:
        page_count = len(pdf.pages)

        # Batch processing for large PDFs to avoid memory issues
        BATCH_SIZE = 20
        for batch_start in range(0, page_count, BATCH_SIZE):
            batch_end = min(batch_start + BATCH_SIZE, page_count)
            batch_pages = pdf.pages[batch_start:batch_end]

            for i, page in enumerate(batch_pages):
                page_num = batch_start + i + 1

                # Extract tables as markdown chunks (with explicit chunk_type='table')
                page_tables = _extract_tables_as_markdown(page, page_num)
                table_chunks.extend(page_tables)

                # Extract regular text
                page_text = page.extract_text() or ""

                # OCR fallback for scanned pages with no extractable text
                if not page_text.strip():
                    try:
                        page_image = page.to_image(resolution=300)
                        import pytesseract
                        page_text = pytesseract.image_to_string(page_image.original) or ""
                        if page_text.strip():
                            logger.info(f"OCR extracted text from scanned page {page_num}")
                    except Exception as ocr_err:
                        logger.debug(f"OCR fallback unavailable for page {page_num}: {ocr_err}")

                # D9: Skip empty pages - avoid creating empty chunks
                if not page_text.strip():
                    logger.debug(f"Skipping empty page {page_num}")
                    continue

                full_text += f"\n\n--- Page {page_num} ---\n\n{page_text}"

                # Try to extract section headers from this page
                headers = re.findall(r'^[A-Z][A-Z\s]{2,}(?::|$)', page_text, re.MULTILINE)
                if headers:
                    toc["pages"].append({
                        "page": page_num,
                        "headers": headers[:5]  # Limit to avoid noise
                    })

            # Clear batch references for memory management
            del batch_pages

    # Return table chunks separately - they already have chunk_type='table' set
    return full_text, page_count, toc, table_chunks


def _extract_pdf_with_pypdf(filepath: str) -> tuple[str, int, Dict, List[Dict]]:
    """
    Fallback: Extract text from PDF using PyPDFLoader.

    Returns:
        tuple: (full_text, page_count, toc, table_chunks)
        - PyPDF doesn't extract tables well, so table_chunks is empty
    """
    from langchain_community.document_loaders import PyPDFLoader

    loader = PyPDFLoader(filepath)
    pages = loader.load()

    full_text = ""
    toc = {"pages": []}

    for i, page in enumerate(pages):
        page_text = page.page_content
        full_text += f"\n\n--- Page {i+1} ---\n\n{page_text}"

        # Try to extract section headers from this page
        headers = re.findall(r'^[A-Z][A-Z\s]{2,}(?::|$)', page_text, re.MULTILINE)
        if headers:
            toc["pages"].append({
                "page": i + 1,
                "headers": headers[:5]  # Limit to avoid noise
            })

    # PyPDF fallback doesn't extract tables - return empty list
    return full_text, len(pages), toc, []


def extract_docx_text(filepath: str) -> tuple[str, int, Dict]:
    """Extract text and structure from DOCX."""
    from docx import Document

    doc = Document(filepath)

    full_text = ""
    toc = {"sections": []}
    current_section = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect headers by style
        if para.style.name.startswith('Heading'):
            current_section = text
            toc["sections"].append(text)
            full_text += f"\n\n## {text}\n\n"
        else:
            full_text += text + "\n"

    # Extract tables
    for table in doc.tables:
        table_text = "\n"
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            table_text += row_text + "\n"
        full_text += table_text + "\n"

    return full_text, len(doc.paragraphs) // 40 + 1, toc  # Rough page estimate


def extract_text_file(filepath: str) -> tuple[str, int, Dict]:
    """Extract text from TXT/MD files."""
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        text = f.read()

    # Try to extract headers
    headers = re.findall(r'^#+\s+(.+)$', text, re.MULTILINE)
    toc = {"headers": headers} if headers else {}

    return text, len(text) // 2000 + 1, toc


# ============================================================================
# FACT EXTRACTION PROMPTS
# ============================================================================

FACT_EXTRACTION_PROMPT = """You are analyzing a technical document chunk for a lift/elevator maintenance system.

Extract ALL useful facts from this content. For each fact, identify:
- fact_type: spec|procedure|warning|wiring|troubleshooting|part|setting|measurement
- content: The actual fact (clear, standalone statement)
- keywords: Key terms for search (3-5 words)

Document info:
- Filename: {filename}
- Lift models: {lift_models}
- Section: {section}
- Page: {page}

Content to analyze:
---
{content}
---

Return a JSON array of facts. Example:
[
  {{
    "fact_type": "spec",
    "content": "Drive motor torque setting is 25Nm for Elvoron models",
    "keywords": ["torque", "motor", "25Nm", "drive"]
  }},
  {{
    "fact_type": "procedure",
    "content": "To reset the safety loop: 1) Turn off main power, 2) Wait 30 seconds, 3) Reset the safety relay on terminal T6",
    "keywords": ["reset", "safety loop", "T6", "relay"]
  }}
]

Extract ALL facts - specs, procedures, warnings, wiring info, part numbers, settings.
If no useful facts, return empty array [].
Return ONLY valid JSON, no explanation."""


QA_GENERATION_PROMPT = """Based on these extracted facts from lift/elevator documentation, generate likely questions a technician would ask.

Facts:
{facts}

Generate 3-5 natural questions that these facts answer. Include:
- Spec lookups ("What is the torque for...")
- How-to questions ("How do I reset...")
- Troubleshooting ("Why is the lift showing error...")

Return JSON array:
[
  {{
    "question": "What is the torque setting for the Elvoron drive motor?",
    "answer_summary": "25Nm",
    "full_answer": "The drive motor torque setting for Elvoron models is 25Nm. This is set during installation using a torque wrench.",
    "category": "spec"
  }}
]

Return ONLY valid JSON."""


ENTITY_EXTRACTION_PROMPT = """Extract specific entities from these lift maintenance facts:

Facts:
{facts}

Extract:
- Part numbers (P/N, part references)
- Specifications (measurements with values and units)
- Error codes (E01, fault codes, etc.)
- Terminal numbers (T1, T6, etc.)
- Settings (parameters, values)

Return JSON:
[
  {{
    "entity_type": "spec",
    "identifier": "25Nm",
    "description": "Drive motor torque setting",
    "value": "25",
    "unit": "Nm"
  }},
  {{
    "entity_type": "terminal",
    "identifier": "T6",
    "description": "Safety relay reset terminal"
  }}
]

Return ONLY valid JSON, no explanation."""


# ============================================================================
# CLAUDE EXTRACTION
# ============================================================================

async def extract_facts_with_claude(chunk_content: str, context: Dict) -> List[Dict]:
    """Use Claude to extract structured facts from a chunk."""
    import httpx

    prompt = FACT_EXTRACTION_PROMPT.format(
        filename=context.get("filename", "unknown"),
        lift_models=", ".join(context.get("lift_models", [])) or "unknown",
        section=context.get("section", "unknown"),
        page=context.get("page", "unknown"),
        content=chunk_content
    )

    # Call Claude CLI or API
    try:
        result = await _call_claude(prompt)
        facts = json.loads(result)
        return facts if isinstance(facts, list) else []
    except Exception as e:
        logger.error(f"Fact extraction failed: {e}")
        return []


async def generate_qa_pairs(facts: List[Dict]) -> List[Dict]:
    """Generate Q&A pairs from extracted facts."""
    if not facts:
        return []

    facts_text = json.dumps(facts, indent=2)
    prompt = QA_GENERATION_PROMPT.format(facts=facts_text)

    try:
        result = await _call_claude(prompt)
        qa_pairs = json.loads(result)
        return qa_pairs if isinstance(qa_pairs, list) else []
    except Exception as e:
        logger.error(f"QA generation failed: {e}")
        return []


async def extract_entities(facts: List[Dict]) -> List[Dict]:
    """Extract entities from facts."""
    if not facts:
        return []

    facts_text = json.dumps(facts, indent=2)
    prompt = ENTITY_EXTRACTION_PROMPT.format(facts=facts_text)

    try:
        result = await _call_claude(prompt)
        entities = json.loads(result)
        return entities if isinstance(entities, list) else []
    except Exception as e:
        logger.error(f"Entity extraction failed: {e}")
        return []


async def _call_claude(prompt: str) -> str:
    """Call Claude CLI with Opus 4.5 for fact extraction."""
    import subprocess
    import asyncio
    import tempfile

    # Use Claude CLI with Opus 4.5
    try:
        # Write prompt to temp file to handle long prompts
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(prompt)
            prompt_file = f.name

        # Run Claude CLI with Opus 4.5 in print mode
        # Read prompt from file using shell redirection
        process = await asyncio.create_subprocess_shell(
            f'cat "{prompt_file}" | claude --model opus --print --dangerously-skip-permissions --output-format text',
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )

        stdout, stderr = await asyncio.wait_for(
            process.communicate(),
            timeout=180  # 3 minute timeout for complex extractions
        )

        # Clean up temp file
        try:
            os.unlink(prompt_file)
        except:
            pass

        if process.returncode == 0:
            result = stdout.decode('utf-8').strip()
            # Try to extract JSON from the response
            json_match = re.search(r'\[[\s\S]*\]', result)
            if json_match:
                return json_match.group(0)
            return result
        else:
            error_msg = stderr.decode('utf-8').strip()
            if error_msg:
                logger.error(f"Claude CLI error: {error_msg}")
            return "[]"

    except asyncio.TimeoutError:
        logger.error("Claude CLI timed out")
        return "[]"
    except FileNotFoundError:
        logger.error("Claude CLI not found. Install with: npm install -g @anthropic-ai/claude-code")
        return "[]"
    except Exception as e:
        logger.error(f"Claude CLI call failed: {e}")
        return "[]"


# ============================================================================
# SAVE TO DATABASE
# ============================================================================

def save_chunks(document_id: int, chunks: List[Dict], lift_models: List[str]) -> int:
    """
    Save chunks to database with embeddings.
    Uses batch processing (32 chunks at a time) for efficient embedding generation.

    Preserves explicitly-set chunk_type (e.g., 'table' from table extraction) and
    section_path metadata. Falls back to detection for chunks without explicit type.

    Returns:
        int: Number of chunks successfully saved to database
    """
    if not chunks:
        logger.warning(f"save_chunks called with empty chunks list for document {document_id}")
        return 0

    conn = get_db()
    cur = conn.cursor()

    BATCH_SIZE = 32
    chunk_data = []

    # Prepare all chunk data with types
    for i, chunk in enumerate(chunks):
        # Use explicit chunk_type if set (e.g., from table extraction), else detect
        chunk_type = chunk.get("chunk_type") or detect_chunk_type(chunk["content"])
        section_path = chunk.get("section_path", "")

        # D10: Get page_start and page_end, falling back to page_number for table chunks
        page_start = chunk.get("page_start") or chunk.get("page_number")
        page_end = chunk.get("page_end") or chunk.get("page_number")

        chunk_data.append({
            "index": i,
            "content": chunk["content"],
            "chunk_type": chunk_type,
            "section_path": section_path,
            "tokens": chunk.get("tokens", 0),
            "page_start": page_start,
            "page_end": page_end
        })

    # Process in batches for efficient embedding generation
    for batch_start in range(0, len(chunk_data), BATCH_SIZE):
        batch = chunk_data[batch_start:batch_start + BATCH_SIZE]
        batch_texts = [c["content"] for c in batch]

        # Generate embeddings in batch
        embeddings = generate_embeddings_batch(batch_texts)

        # Insert each chunk with its embedding
        for j, chunk_item in enumerate(batch):
            embedding = embeddings[j]
            embedding_str = None
            if embedding is not None:
                # Format embedding as PostgreSQL vector string
                embedding_str = "[" + ",".join(str(x) for x in embedding) + "]"

            cur.execute("""
                INSERT INTO chunks (document_id, chunk_index, content, chunk_type, section_path,
                                   page_start, page_end, token_count, lift_models, embedding)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s::vector)
                RETURNING id
            """, (
                document_id,
                chunk_item["index"],
                chunk_item["content"],
                chunk_item["chunk_type"],
                chunk_item["section_path"] or None,
                chunk_item["page_start"],
                chunk_item["page_end"],
                chunk_item["tokens"],
                lift_models,
                embedding_str
            ))

    conn.commit()

    # Verify chunks were actually saved
    cur.execute("SELECT COUNT(*) FROM chunks WHERE document_id = %s", (document_id,))
    saved_count = cur.fetchone()[0]

    cur.close()
    conn.close()

    if saved_count == 0:
        logger.error(f"CRITICAL: 0 chunks saved for document {document_id} despite {len(chunks)} chunks provided")
    else:
        logger.info(f"Saved {saved_count} chunks with embeddings for document {document_id}")

    return saved_count


def save_facts(document_id: int, chunk_id: int, facts: List[Dict], context: Dict):
    """
    Save extracted facts to database with embeddings.
    Uses batch processing (32 facts at a time) for efficient embedding generation.
    """
    if not facts:
        return

    conn = get_db()
    cur = conn.cursor()

    BATCH_SIZE = 32

    # Process in batches for efficient embedding generation
    for batch_start in range(0, len(facts), BATCH_SIZE):
        batch = facts[batch_start:batch_start + BATCH_SIZE]
        batch_texts = [f.get("content", "") for f in batch]

        # Generate embeddings in batch
        embeddings = generate_embeddings_batch(batch_texts)

        # Insert each fact with its embedding
        for j, fact in enumerate(batch):
            embedding = embeddings[j]
            embedding_str = None
            if embedding is not None:
                # Format embedding as PostgreSQL vector string
                embedding_str = "[" + ",".join(str(x) for x in embedding) + "]"

            cur.execute("""
                INSERT INTO facts (document_id, chunk_id, fact_type, lift_models, content, keywords, page, section_path, embedding)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s::vector)
            """, (
                document_id,
                chunk_id,
                fact.get("fact_type", "general"),
                context.get("lift_models", []),
                fact.get("content", ""),
                fact.get("keywords", []),
                context.get("page"),
                context.get("section"),
                embedding_str
            ))

    conn.commit()
    cur.close()
    conn.close()
    logger.info(f"Saved {len(facts)} facts with embeddings for document {document_id}")


def save_qa_pairs(qa_pairs: List[Dict], lift_models: List[str], fact_ids: List[int]):
    """Save Q&A pairs to database."""
    conn = get_db()
    cur = conn.cursor()

    for qa in qa_pairs:
        cur.execute("""
            INSERT INTO qa_pairs (question, answer_summary, full_answer, fact_ids, lift_models, category)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (
            qa.get("question", ""),
            qa.get("answer_summary", ""),
            qa.get("full_answer", ""),
            fact_ids,
            lift_models,
            qa.get("category", "general")
        ))

    conn.commit()
    cur.close()
    conn.close()


def save_entities(entities: List[Dict], lift_models: List[str], fact_ids: List[int]):
    """Save entities to database."""
    conn = get_db()
    cur = conn.cursor()

    for entity in entities:
        identifier = entity.get("identifier", "")
        cur.execute("""
            INSERT INTO entities (entity_type, identifier, identifier_normalized, description, value, unit, lift_models, fact_ids)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            ON CONFLICT (entity_type, identifier, lift_models) DO UPDATE
            SET description = EXCLUDED.description, fact_ids = EXCLUDED.fact_ids
        """, (
            entity.get("entity_type", "unknown"),
            identifier,
            identifier.lower().replace(" ", ""),
            entity.get("description", ""),
            entity.get("value"),
            entity.get("unit"),
            lift_models,
            fact_ids
        ))

    conn.commit()
    cur.close()
    conn.close()


def update_document_status(document_id: int, status: str, **kwargs):
    """Update document indexing status."""
    conn = get_db()
    cur = conn.cursor()

    updates = ["index_status = %s", "updated_at = NOW()"]
    values = [status]

    if status == 'complete':
        updates.append("indexed_at = NOW()")

    for key, value in kwargs.items():
        if key in ['title', 'summary', 'page_count', 'error_message']:
            updates.append(f"{key} = %s")
            values.append(value)
        elif key == 'toc':
            updates.append("toc = %s")
            values.append(Json(value))
        elif key == 'lift_models':
            updates.append("lift_models = %s")
            values.append(value)

    values.append(document_id)

    cur.execute(f"""
        UPDATE documents SET {', '.join(updates)} WHERE id = %s
    """, values)

    conn.commit()
    cur.close()
    conn.close()


# ============================================================================
# MAIN INDEXING PIPELINE
# ============================================================================

async def index_document(filepath: str, lift_model_hint: str = None) -> Dict:
    """
    Full indexing pipeline for a single document.

    1. Register document
    2. Extract text and structure
    3. Smart chunk
    4. Extract facts using Claude
    5. Generate Q&A pairs
    6. Extract entities
    """
    result = {
        "filepath": filepath,
        "status": "pending",
        "chunks": 0,
        "facts": 0,
        "qa_pairs": 0,
        "entities": 0,
        "errors": []
    }

    try:
        # 1. Register document
        doc_id = register_document(filepath, lift_model_hint)
        result["document_id"] = doc_id

        update_document_status(doc_id, 'processing')

        # 2. Extract text based on file type
        file_type = get_file_type(filepath)
        filename = os.path.basename(filepath)
        table_chunks = []  # Pre-extracted table chunks with explicit chunk_type

        if file_type == 'pdf':
            text, page_count, toc, table_chunks = extract_pdf_text(filepath)
        elif file_type == 'document':
            text, page_count, toc = extract_docx_text(filepath)
        elif file_type == 'text':
            text, page_count, toc = extract_text_file(filepath)
        elif file_type == 'image':
            # Images handled separately
            result["status"] = "skipped"
            result["message"] = "Image files indexed separately"
            return result
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        # Validate text extraction - catch empty documents early
        if not text or not text.strip():
            logger.warning(f"Empty text extraction for {filepath} - may be scanned/image-only PDF")
            update_document_status(
                doc_id, 'error',
                error_message="No text could be extracted - document may be scanned/image-only"
            )
            result["status"] = "error"
            result["errors"].append("Empty text extraction - document may be scanned/image-only")
            return result

        # Detect lift models from content
        lift_models = detect_lift_models(text, filename)
        if lift_model_hint and lift_model_hint not in lift_models:
            lift_models.insert(0, lift_model_hint)

        # 3. Smart chunk text content
        text_chunks = smart_chunk_text(text)

        # Merge text chunks with pre-extracted table chunks (which have explicit chunk_type='table')
        # Table chunks are inserted at appropriate positions or appended at end
        all_chunks = text_chunks + table_chunks  # Tables at end for now

        # Validate we have chunks to save
        if not all_chunks:
            logger.warning(f"Chunking produced 0 chunks for {filepath}")
            update_document_status(
                doc_id, 'error',
                error_message="Text extraction succeeded but chunking produced 0 chunks"
            )
            result["status"] = "error"
            result["errors"].append("Chunking produced 0 chunks")
            return result

        # Save chunks and verify
        saved_count = save_chunks(doc_id, all_chunks, lift_models)
        result["chunks"] = saved_count

        # CRITICAL: Verify at least 1 chunk was saved - prevents silent failures
        if saved_count == 0:
            logger.error(f"CRITICAL: 0 chunks saved for {filepath} despite {len(all_chunks)} chunks generated")
            update_document_status(
                doc_id, 'error',
                error_message=f"Failed to save chunks: {len(all_chunks)} chunks generated but 0 saved to database"
            )
            result["status"] = "error"
            result["errors"].append(f"Failed to save chunks to database")
            return result

        # 4-6. Process each chunk with Claude
        all_facts = []

        for i, chunk in enumerate(all_chunks):
            # Use section_path if available, otherwise generic chunk identifier
            section = chunk.get("section_path") or f"Chunk {i+1}"
            page = chunk.get("page_number") or (i + 1)

            context = {
                "filename": filename,
                "lift_models": lift_models,
                "section": section,
                "page": page
            }

            # Extract facts
            facts = await extract_facts_with_claude(chunk["content"], context)
            if facts:
                save_facts(doc_id, None, facts, context)  # chunk_id linkage done later
                all_facts.extend(facts)

        result["facts"] = len(all_facts)

        # Generate Q&A pairs from all facts
        if all_facts:
            qa_pairs = await generate_qa_pairs(all_facts)
            if qa_pairs:
                save_qa_pairs(qa_pairs, lift_models, [])
                result["qa_pairs"] = len(qa_pairs)

            # Extract entities
            entities = await extract_entities(all_facts)
            if entities:
                save_entities(entities, lift_models, [])
                result["entities"] = len(entities)

        # Update document with metadata
        update_document_status(
            doc_id, 'complete',
            page_count=page_count,
            toc=toc,
            lift_models=lift_models
        )

        result["status"] = "complete"

    except Exception as e:
        logger.error(f"Error indexing {filepath}: {e}")
        result["status"] = "error"
        result["errors"].append(str(e))

        if "doc_id" in result:
            update_document_status(result["doc_id"], 'error', error_message=str(e))

    return result


def index_document_sync(filepath: str, lift_model_hint: str = None) -> Dict:
    """Synchronous wrapper for index_document."""
    import asyncio
    return asyncio.run(index_document(filepath, lift_model_hint))
